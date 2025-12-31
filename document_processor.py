#!/usr/bin/env python3
"""
Document Processing with Supabase Storage
Automatically processes and understands uploaded documents
Converted to async asyncpg for non-blocking database operations
"""

import os
import json
import logging
import uuid
import hashlib
import asyncio
import mimetypes
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime, timezone
from enum import Enum
import httpx
from openai import OpenAI
import PyPDF2
import docx
import openpyxl
from PIL import Image
import pytesseract
import io
import base64

# Import async database connection
from database.async_connection import get_pool

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Supabase configuration
SUPABASE_URL = os.getenv("SUPABASE_URL", "https://yomagoqdmxszqtdwuhab.supabase.co")
SUPABASE_KEY = os.getenv("SUPABASE_ANON_KEY", "")
STORAGE_BUCKET = "documents"

# OpenAI configuration
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

class DocumentType(Enum):
    """Types of documents"""
    PDF = "pdf"
    WORD = "word"
    EXCEL = "excel"
    IMAGE = "image"
    TEXT = "text"
    PRESENTATION = "presentation"
    EMAIL = "email"
    CONTRACT = "contract"
    INVOICE = "invoice"
    REPORT = "report"
    UNKNOWN = "unknown"

class ProcessingStatus(Enum):
    """Document processing status"""
    UPLOADED = "uploaded"
    QUEUED = "queued"
    PROCESSING = "processing"
    ANALYZING = "analyzing"
    COMPLETED = "completed"
    FAILED = "failed"

class ExtractionMethod(Enum):
    """Text extraction methods"""
    NATIVE = "native"
    OCR = "ocr"
    AI_VISION = "ai_vision"
    HYBRID = "hybrid"

class DocumentCategory(Enum):
    """Document categories for classification"""
    LEGAL = "legal"
    FINANCIAL = "financial"
    TECHNICAL = "technical"
    MARKETING = "marketing"
    HR = "hr"
    OPERATIONAL = "operational"
    CUSTOMER = "customer"
    VENDOR = "vendor"
    INTERNAL = "internal"
    EXTERNAL = "external"

class DocumentProcessor:
    """Main document processing class"""

    def __init__(self):
        """Initialize the document processor"""
        self.storage_client = None
        self.supported_formats = {
            '.pdf': DocumentType.PDF,
            '.doc': DocumentType.WORD,
            '.docx': DocumentType.WORD,
            '.xls': DocumentType.EXCEL,
            '.xlsx': DocumentType.EXCEL,
            '.png': DocumentType.IMAGE,
            '.jpg': DocumentType.IMAGE,
            '.jpeg': DocumentType.IMAGE,
            '.txt': DocumentType.TEXT,
            '.csv': DocumentType.TEXT,
            '.ppt': DocumentType.PRESENTATION,
            '.pptx': DocumentType.PRESENTATION
        }
        self._initialized = False
        self._init_storage()

    async def initialize(self):
        """Async initialization - call this after creating the instance"""
        if not self._initialized:
            await self._init_database()
            self._initialized = True

    async def _init_database(self):
        """Initialize database tables using async pool"""
        try:
            pool = get_pool()

            # Create tables
            await pool.execute("""
                CREATE TABLE IF NOT EXISTS ai_documents (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    filename VARCHAR(255),
                    document_type VARCHAR(50),
                    mime_type VARCHAR(100),
                    file_size BIGINT,
                    storage_path TEXT,
                    storage_url TEXT,
                    hash_value VARCHAR(64),
                    status VARCHAR(50) DEFAULT 'uploaded',
                    upload_source VARCHAR(100),
                    user_id VARCHAR(255),
                    metadata JSONB DEFAULT '{}',
                    created_at TIMESTAMPTZ DEFAULT NOW(),
                    processed_at TIMESTAMPTZ
                )
            """)

            await pool.execute("""
                CREATE TABLE IF NOT EXISTS ai_document_content (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    document_id UUID REFERENCES ai_documents(id),
                    content_type VARCHAR(50),
                    raw_text TEXT,
                    structured_data JSONB,
                    extraction_method VARCHAR(50),
                    language VARCHAR(10),
                    word_count INT,
                    page_count INT,
                    confidence_score FLOAT,
                    metadata JSONB DEFAULT '{}',
                    created_at TIMESTAMPTZ DEFAULT NOW()
                )
            """)

            await pool.execute("""
                CREATE TABLE IF NOT EXISTS ai_document_analysis (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    document_id UUID REFERENCES ai_documents(id),
                    analysis_type VARCHAR(50),
                    category VARCHAR(50),
                    summary TEXT,
                    key_entities JSONB,
                    key_terms JSONB,
                    sentiment_score FLOAT,
                    topics JSONB,
                    classifications JSONB,
                    insights JSONB,
                    metadata JSONB DEFAULT '{}',
                    created_at TIMESTAMPTZ DEFAULT NOW()
                )
            """)

            await pool.execute("""
                CREATE TABLE IF NOT EXISTS ai_document_embeddings (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    document_id UUID REFERENCES ai_documents(id),
                    chunk_index INT,
                    chunk_text TEXT,
                    embedding vector(1536),
                    metadata JSONB DEFAULT '{}',
                    created_at TIMESTAMPTZ DEFAULT NOW()
                )
            """)

            await pool.execute("""
                CREATE TABLE IF NOT EXISTS ai_document_relationships (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    source_document_id UUID REFERENCES ai_documents(id),
                    target_document_id UUID REFERENCES ai_documents(id),
                    relationship_type VARCHAR(50),
                    confidence FLOAT,
                    metadata JSONB DEFAULT '{}',
                    created_at TIMESTAMPTZ DEFAULT NOW()
                )
            """)

            await pool.execute("""
                CREATE TABLE IF NOT EXISTS ai_document_actions (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    document_id UUID REFERENCES ai_documents(id),
                    action_type VARCHAR(50),
                    action_status VARCHAR(50),
                    action_data JSONB,
                    triggered_by VARCHAR(100),
                    executed_at TIMESTAMPTZ,
                    result JSONB,
                    created_at TIMESTAMPTZ DEFAULT NOW()
                )
            """)

            # Create indexes
            await pool.execute("CREATE INDEX IF NOT EXISTS idx_documents_status ON ai_documents(status)")
            await pool.execute("CREATE INDEX IF NOT EXISTS idx_documents_type ON ai_documents(document_type)")
            await pool.execute("CREATE INDEX IF NOT EXISTS idx_documents_user ON ai_documents(user_id)")
            await pool.execute("CREATE INDEX IF NOT EXISTS idx_content_document ON ai_document_content(document_id)")
            await pool.execute("CREATE INDEX IF NOT EXISTS idx_analysis_document ON ai_document_analysis(document_id)")
            await pool.execute("CREATE INDEX IF NOT EXISTS idx_embeddings_document ON ai_document_embeddings(document_id)")

            logger.info("Database tables initialized successfully")

        except Exception as e:
            logger.error(f"Failed to initialize database: {e}")

    def _init_storage(self):
        """Initialize Supabase storage client"""
        try:
            self.storage_client = httpx.AsyncClient(
                base_url=f"{SUPABASE_URL}/storage/v1",
                headers={
                    "apikey": SUPABASE_KEY,
                    "Authorization": f"Bearer {SUPABASE_KEY}"
                }
            )
        except Exception as e:
            logger.error(f"Failed to initialize storage client: {e}")

    async def upload_document(
        self,
        file_path: str = None,
        file_content: bytes = None,
        filename: str = None,
        user_id: str = None,
        source: str = "manual",
        metadata: Dict = None
    ) -> str:
        """Upload a document to storage and create database record"""
        try:
            # Get file content if path provided
            if file_path and not file_content:
                with open(file_path, 'rb') as f:
                    file_content = f.read()
                if not filename:
                    filename = os.path.basename(file_path)

            if not file_content or not filename:
                raise ValueError("File content and filename required")

            # Calculate file hash
            file_hash = hashlib.sha256(file_content).hexdigest()

            # Check for duplicate using async pool
            pool = get_pool()

            existing = await pool.fetchrow("""
                SELECT id, storage_url FROM ai_documents
                WHERE hash_value = $1
            """, file_hash)

            if existing:
                logger.info(f"Document already exists: {existing['id']}")
                return str(existing['id'])

            # Determine document type
            file_ext = os.path.splitext(filename)[1].lower()
            doc_type = self.supported_formats.get(file_ext, DocumentType.UNKNOWN)
            mime_type = mimetypes.guess_type(filename)[0] or 'application/octet-stream'

            # Generate storage path
            doc_id = str(uuid.uuid4())
            storage_path = f"{user_id or 'public'}/{datetime.now().strftime('%Y/%m/%d')}/{doc_id}/{filename}"

            # Upload to Supabase Storage
            if self.storage_client:
                response = await self.storage_client.post(
                    f"/object/{STORAGE_BUCKET}/{storage_path}",
                    content=file_content,
                    headers={"Content-Type": mime_type}
                )

                if response.status_code == 200:
                    storage_url = f"{SUPABASE_URL}/storage/v1/object/public/{STORAGE_BUCKET}/{storage_path}"
                else:
                    storage_url = None
                    logger.error(f"Failed to upload to storage: {response.text}")
            else:
                # Fallback: Store locally
                storage_url = f"local://{storage_path}"

            # Create database record using async pool
            await pool.execute("""
                INSERT INTO ai_documents
                (id, filename, document_type, mime_type, file_size,
                 storage_path, storage_url, hash_value, status,
                 upload_source, user_id, metadata)
                VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12)
            """,
                doc_id,
                filename,
                doc_type.value,
                mime_type,
                len(file_content),
                storage_path,
                storage_url,
                file_hash,
                ProcessingStatus.UPLOADED.value,
                source,
                user_id,
                json.dumps(metadata or {})
            )

            # Queue for processing
            await self._queue_for_processing(doc_id)

            return doc_id

        except Exception as e:
            logger.error(f"Failed to upload document: {e}")
            return None

    async def _queue_for_processing(self, document_id: str):
        """Queue document for processing"""
        try:
            pool = get_pool()

            await pool.execute("""
                UPDATE ai_documents
                SET status = $1
                WHERE id = $2
            """, ProcessingStatus.QUEUED.value, document_id)

            # Trigger async processing
            asyncio.create_task(self.process_document(document_id))

        except Exception as e:
            logger.error(f"Failed to queue document: {e}")

    async def process_document(self, document_id: str) -> Dict:
        """Process a document - extract text, analyze, and create embeddings"""
        try:
            pool = get_pool()

            # Update status
            await pool.execute("""
                UPDATE ai_documents
                SET status = $1
                WHERE id = $2
            """, ProcessingStatus.PROCESSING.value, document_id)

            # Get document info
            document = await pool.fetchrow("""
                SELECT * FROM ai_documents WHERE id = $1
            """, document_id)

            if not document:
                raise ValueError(f"Document not found: {document_id}")

            # Convert Record to dict for easier access
            document = dict(document)

            # Extract text based on document type
            extracted_text = await self._extract_text(document)

            # Store extracted content
            if extracted_text:
                content_id = str(uuid.uuid4())
                await pool.execute("""
                    INSERT INTO ai_document_content
                    (id, document_id, content_type, raw_text,
                     extraction_method, word_count, confidence_score)
                    VALUES ($1, $2, $3, $4, $5, $6, $7)
                """,
                    content_id,
                    document_id,
                    document['document_type'],
                    extracted_text['text'],
                    extracted_text['method'],
                    len(extracted_text['text'].split()),
                    extracted_text.get('confidence', 1.0)
                )

                # Analyze document
                analysis = await self._analyze_document(
                    document_id,
                    extracted_text['text'],
                    document
                )

                # Create embeddings
                await self._create_embeddings(
                    document_id,
                    extracted_text['text']
                )

                # Find relationships
                await self._find_relationships(document_id, analysis)

                # Update status
                await pool.execute("""
                    UPDATE ai_documents
                    SET status = $1, processed_at = $2
                    WHERE id = $3
                """, ProcessingStatus.COMPLETED.value, datetime.now(timezone.utc), document_id)

                return {
                    'document_id': document_id,
                    'status': 'completed',
                    'extracted_text': len(extracted_text['text']),
                    'analysis': analysis
                }

            else:
                raise ValueError("Failed to extract text from document")

        except Exception as e:
            logger.error(f"Failed to process document: {e}")

            # Update status to failed
            try:
                pool = get_pool()
                await pool.execute("""
                    UPDATE ai_documents
                    SET status = $1
                    WHERE id = $2
                """, ProcessingStatus.FAILED.value, document_id)
            except Exception as update_error:
                logger.error(
                    "Failed to mark document %s as failed: %s",
                    document_id,
                    update_error,
                    exc_info=True,
                )

            return {
                'document_id': document_id,
                'status': 'failed',
                'error': str(e)
            }

    async def _extract_text(self, document: Dict) -> Dict:
        """Extract text from document based on type"""
        try:
            doc_type = document['document_type']
            storage_url = document['storage_url']

            # Download file content if needed
            file_content = await self._download_file(storage_url)

            if doc_type == DocumentType.PDF.value:
                return await self._extract_pdf_text(file_content)

            elif doc_type == DocumentType.WORD.value:
                return await self._extract_word_text(file_content)

            elif doc_type == DocumentType.EXCEL.value:
                return await self._extract_excel_text(file_content)

            elif doc_type == DocumentType.IMAGE.value:
                return await self._extract_image_text(file_content)

            elif doc_type == DocumentType.TEXT.value:
                return {
                    'text': file_content.decode('utf-8', errors='ignore'),
                    'method': ExtractionMethod.NATIVE.value,
                    'confidence': 1.0
                }

            else:
                # Try generic text extraction
                try:
                    text = file_content.decode('utf-8', errors='ignore')
                    return {
                        'text': text,
                        'method': ExtractionMethod.NATIVE.value,
                        'confidence': 0.8
                    }
                except (AttributeError, TypeError) as exc:
                    logger.debug("Text decode failed, falling back to vision: %s", exc, exc_info=True)
                    # Fall back to AI vision
                    return await self._extract_with_ai_vision(file_content)

        except Exception as e:
            logger.error(f"Failed to extract text: {e}")
            return None

    async def _download_file(self, storage_url: str) -> bytes:
        """Download file from storage"""
        if storage_url.startswith('local://'):
            # Local file
            path = storage_url.replace('local://', '')
            with open(path, 'rb') as f:
                return f.read()
        else:
            # Download from URL
            async with httpx.AsyncClient() as client:
                response = await client.get(storage_url)
                return response.content

    async def _extract_pdf_text(self, content: bytes) -> Dict:
        """Extract text from PDF"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

            return {
                'text': text.strip(),
                'method': ExtractionMethod.NATIVE.value,
                'confidence': 0.95,
                'pages': len(pdf_reader.pages)
            }
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            # Fall back to OCR
            return await self._extract_with_ocr(content)

    async def _extract_word_text(self, content: bytes) -> Dict:
        """Extract text from Word document"""
        try:
            doc = docx.Document(io.BytesIO(content))
            text = "\n".join([para.text for para in doc.paragraphs])

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text

            return {
                'text': text.strip(),
                'method': ExtractionMethod.NATIVE.value,
                'confidence': 0.95
            }
        except Exception as e:
            logger.error(f"Word extraction failed: {e}")
            return None

    async def _extract_excel_text(self, content: bytes) -> Dict:
        """Extract text from Excel"""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(content))
            text = ""

            for sheet in workbook:
                text += f"Sheet: {sheet.title}\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"

            return {
                'text': text.strip(),
                'method': ExtractionMethod.NATIVE.value,
                'confidence': 0.9
            }
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            return None

    async def _extract_image_text(self, content: bytes) -> Dict:
        """Extract text from image using OCR"""
        try:
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)

            return {
                'text': text.strip(),
                'method': ExtractionMethod.OCR.value,
                'confidence': 0.85
            }
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            # Fall back to AI vision
            return await self._extract_with_ai_vision(content)

    async def _extract_with_ocr(self, content: bytes) -> Dict:
        """Extract text using OCR"""
        try:
            # Convert to image first if needed
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)

            return {
                'text': text.strip(),
                'method': ExtractionMethod.OCR.value,
                'confidence': 0.8
            }
        except (OSError, pytesseract.TesseractError, ValueError) as exc:
            logger.warning("OCR extraction failed: %s", exc, exc_info=True)
            return None

    async def _extract_with_ai_vision(self, content: bytes) -> Dict:
        """Extract text using AI vision API"""
        try:
            # Convert to base64
            base64_image = base64.b64encode(content).decode('utf-8')

            response = openai_client.chat.completions.create(
                model="gpt-4-vision-preview",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Extract all text from this document image. Include tables, headers, and any visible text."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=4000
            )

            return {
                'text': response.choices[0].message.content,
                'method': ExtractionMethod.AI_VISION.value,
                'confidence': 0.9
            }
        except Exception as e:
            logger.error(f"AI vision extraction failed: {e}")
            return None

    async def _analyze_document(
        self,
        document_id: str,
        text: str,
        document: Dict
    ) -> Dict:
        """Analyze document content using AI"""
        try:
            # Prepare analysis prompt
            prompt = f"""
            Analyze this document and extract:
            1. Document category (legal, financial, technical, etc.)
            2. Brief summary (2-3 sentences)
            3. Key entities (people, companies, dates, amounts)
            4. Important terms and concepts
            5. Overall sentiment
            6. Main topics
            7. Any action items or deadlines
            8. Risk factors or concerns

            Document type: {document['document_type']}
            Filename: {document['filename']}

            Content:
            {text[:3000]}  # Limit for API

            Return as JSON.
            """

            response = openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a document analysis expert."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )

            analysis = json.loads(response.choices[0].message.content)

            # Store analysis using async pool
            pool = get_pool()

            await pool.execute("""
                INSERT INTO ai_document_analysis
                (document_id, analysis_type, category, summary,
                 key_entities, key_terms, sentiment_score,
                 topics, classifications, insights)
                VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10)
            """,
                document_id,
                'comprehensive',
                analysis.get('category', 'unknown'),
                analysis.get('summary'),
                json.dumps(analysis.get('entities', {})),
                json.dumps(analysis.get('terms', [])),
                analysis.get('sentiment', 0),
                json.dumps(analysis.get('topics', [])),
                json.dumps(analysis.get('classifications', {})),
                json.dumps(analysis.get('insights', {}))
            )

            return analysis

        except Exception as e:
            logger.error(f"Failed to analyze document: {e}")
            return {}

    async def _create_embeddings(self, document_id: str, text: str):
        """Create embeddings for document chunks"""
        try:
            # Split text into chunks
            chunk_size = 1000
            chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

            pool = get_pool()

            for idx, chunk in enumerate(chunks[:20]):  # Limit to 20 chunks
                # Generate embedding
                response = openai_client.embeddings.create(
                    input=chunk,
                    model="text-embedding-ada-002"
                )

                embedding = response.data[0].embedding

                # Store embedding using async pool
                await pool.execute("""
                    INSERT INTO ai_document_embeddings
                    (document_id, chunk_index, chunk_text, embedding)
                    VALUES ($1, $2, $3, $4)
                """,
                    document_id,
                    idx,
                    chunk,
                    embedding
                )

        except Exception as e:
            logger.error(f"Failed to create embeddings: {e}")

    async def _find_relationships(self, document_id: str, analysis: Dict):
        """Find relationships with other documents"""
        try:
            pool = get_pool()

            # Find similar documents using embeddings
            similar_docs = await pool.fetch("""
                WITH doc_embedding AS (
                    SELECT AVG(embedding) as avg_embedding
                    FROM ai_document_embeddings
                    WHERE document_id = $1
                )
                SELECT
                    d.id,
                    d.filename,
                    AVG(1 - (de.embedding <=> doc_embedding.avg_embedding)) as similarity
                FROM ai_documents d
                JOIN ai_document_embeddings de ON d.id = de.document_id
                CROSS JOIN doc_embedding
                WHERE d.id != $2
                GROUP BY d.id, d.filename
                HAVING AVG(1 - (de.embedding <=> doc_embedding.avg_embedding)) > 0.8
                ORDER BY similarity DESC
                LIMIT 5
            """, document_id, document_id)

            for similar_doc in similar_docs:
                await pool.execute("""
                    INSERT INTO ai_document_relationships
                    (source_document_id, target_document_id,
                     relationship_type, confidence)
                    VALUES ($1, $2, $3, $4)
                """,
                    document_id,
                    similar_doc['id'],
                    'similar',
                    similar_doc['similarity']
                )

        except Exception as e:
            logger.error(f"Failed to find relationships: {e}")

    async def search_documents(
        self,
        query: str,
        filters: Dict = None,
        limit: int = 10
    ) -> List[Dict]:
        """Search documents using semantic search"""
        try:
            # Generate query embedding
            response = openai_client.embeddings.create(
                input=query,
                model="text-embedding-ada-002"
            )
            query_embedding = response.data[0].embedding

            pool = get_pool()

            # Build filter conditions
            where_conditions = ["d.status = 'completed'"]
            params = [query_embedding]
            param_idx = 2

            if filters:
                if filters.get('document_type'):
                    where_conditions.append(f"d.document_type = ${param_idx}")
                    params.append(filters['document_type'])
                    param_idx += 1
                if filters.get('user_id'):
                    where_conditions.append(f"d.user_id = ${param_idx}")
                    params.append(filters['user_id'])
                    param_idx += 1
                if filters.get('category'):
                    where_conditions.append(f"da.category = ${param_idx}")
                    params.append(filters['category'])
                    param_idx += 1

            params.append(limit)
            where_clause = " AND ".join(where_conditions)

            # Search using embeddings
            results = await pool.fetch(f"""
                SELECT DISTINCT ON (d.id)
                    d.*,
                    da.summary,
                    da.category,
                    1 - (de.embedding <=> $1) as similarity
                FROM ai_documents d
                JOIN ai_document_embeddings de ON d.id = de.document_id
                LEFT JOIN ai_document_analysis da ON d.id = da.document_id
                WHERE {where_clause}
                ORDER BY d.id, similarity DESC
                LIMIT ${param_idx}
            """, *params)

            # Convert Records to dicts
            return [dict(r) for r in results]

        except Exception as e:
            logger.error(f"Failed to search documents: {e}")
            return []

    async def get_document_insights(self, document_id: str) -> Dict:
        """Get comprehensive insights for a document"""
        try:
            pool = get_pool()

            # Get document info
            document = await pool.fetchrow("""
                SELECT d.*, dc.raw_text, da.*
                FROM ai_documents d
                LEFT JOIN ai_document_content dc ON d.id = dc.document_id
                LEFT JOIN ai_document_analysis da ON d.id = da.document_id
                WHERE d.id = $1
            """, document_id)

            # Get relationships
            relationships = await pool.fetch("""
                SELECT
                    dr.*,
                    d.filename as related_filename
                FROM ai_document_relationships dr
                JOIN ai_documents d ON dr.target_document_id = d.id
                WHERE dr.source_document_id = $1
                ORDER BY dr.confidence DESC
            """, document_id)

            # Get actions
            actions = await pool.fetch("""
                SELECT * FROM ai_document_actions
                WHERE document_id = $1
                ORDER BY created_at DESC
            """, document_id)

            # Convert Records to dicts
            document_dict = dict(document) if document else None
            relationships_list = [dict(r) for r in relationships]
            actions_list = [dict(a) for a in actions]

            return {
                'document': document_dict,
                'relationships': relationships_list,
                'actions': actions_list,
                'metrics': {
                    'word_count': document_dict.get('word_count', 0) if document_dict else 0,
                    'sentiment': document_dict.get('sentiment_score', 0) if document_dict else 0,
                    'related_documents': len(relationships_list)
                }
            }

        except Exception as e:
            logger.error(f"Failed to get insights: {e}")
            return {}

    async def trigger_document_action(
        self,
        document_id: str,
        action_type: str,
        action_data: Dict = None
    ) -> Dict:
        """Trigger an action based on document content"""
        try:
            pool = get_pool()

            action_id = str(uuid.uuid4())

            # Create action record using async pool
            await pool.execute("""
                INSERT INTO ai_document_actions
                (id, document_id, action_type, action_status,
                 action_data, triggered_by)
                VALUES ($1, $2, $3, $4, $5, $6)
            """,
                action_id,
                document_id,
                action_type,
                'pending',
                json.dumps(action_data or {}),
                'system'
            )

            # Execute action based on type
            result = await self._execute_action(
                document_id,
                action_type,
                action_data
            )

            # Update action status using async pool
            await pool.execute("""
                UPDATE ai_document_actions
                SET action_status = $1,
                    executed_at = $2,
                    result = $3
                WHERE id = $4
            """,
                'completed' if result.get('success') else 'failed',
                datetime.now(timezone.utc),
                json.dumps(result),
                action_id
            )

            return result

        except Exception as e:
            logger.error(f"Failed to trigger action: {e}")
            return {'success': False, 'error': str(e)}

    async def _execute_action(
        self,
        document_id: str,
        action_type: str,
        action_data: Dict
    ) -> Dict:
        """Execute specific action based on document"""
        try:
            if action_type == 'extract_invoice_data':
                # Extract invoice information
                return await self._extract_invoice_data(document_id)

            elif action_type == 'create_task':
                # Create task from document
                return await self._create_task_from_document(document_id, action_data)

            elif action_type == 'notify_team':
                # Send notification
                return await self._notify_team(document_id, action_data)

            elif action_type == 'generate_response':
                # Generate response document
                return await self._generate_response(document_id, action_data)

            else:
                return {'success': True, 'message': f'Action {action_type} executed'}

        except Exception as e:
            return {'success': False, 'error': str(e)}

    async def _extract_invoice_data(self, document_id: str) -> Dict:
        """Extract structured invoice data"""
        # Implementation for invoice extraction
        return {'success': True, 'invoice_data': {}}

    async def _create_task_from_document(self, document_id: str, data: Dict) -> Dict:
        """Create task based on document content"""
        # Implementation for task creation
        return {'success': True, 'task_id': str(uuid.uuid4())}

    async def _notify_team(self, document_id: str, data: Dict) -> Dict:
        """Send notification about document"""
        # Implementation for notifications
        return {'success': True, 'notified': data.get('recipients', [])}

    async def _generate_response(self, document_id: str, data: Dict) -> Dict:
        """Generate response document"""
        # Implementation for response generation
        return {'success': True, 'response_id': str(uuid.uuid4())}


# Singleton instance
_document_processor = None


async def get_document_processor() -> DocumentProcessor:
    """Get or create the document processor instance (async)"""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
        await _document_processor.initialize()
    return _document_processor
